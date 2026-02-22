"""
Document Processing Module for Ultra Doc-Intelligence.
Handles parsing of PDF, DOCX, and TXT files with intelligent text extraction.
"""

import re
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from enum import Enum
import logging

# Document parsing libraries
import pypdf
import pdfplumber
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    UNKNOWN = "unknown"


@dataclass
class TextChunk:
    """Represents a chunk of text from a document."""
    content: str
    chunk_id: str
    document_id: str
    chunk_index: int
    start_char: int
    end_char: int
    page_number: Optional[int] = None
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert chunk to dictionary for storage."""
        return {
            "content": self.content,
            "chunk_id": self.chunk_id,
            "document_id": self.document_id,
            "chunk_index": self.chunk_index,
            "start_char": self.start_char,
            "end_char": self.end_char,
            "page_number": self.page_number,
            "metadata": self.metadata
        }


@dataclass
class ParsedDocument:
    """Represents a fully parsed document."""
    document_id: str
    filename: str
    document_type: DocumentType
    raw_text: str
    chunks: List[TextChunk]
    page_count: int
    char_count: int
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert document to dictionary."""
        return {
            "document_id": self.document_id,
            "filename": self.filename,
            "document_type": self.document_type.value,
            "page_count": self.page_count,
            "char_count": self.char_count,
            "chunk_count": len(self.chunks),
            "metadata": self.metadata
        }


class DocumentParser:
    """
    Handles parsing of various document formats.
    Extracts text while preserving structure where possible.
    """
    
    def __init__(self):
        self.supported_extensions = {".pdf", ".docx", ".txt"}
    
    def get_document_type(self, filepath: Path) -> DocumentType:
        """Determine document type from file extension."""
        ext = filepath.suffix.lower()
        type_map = {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".txt": DocumentType.TXT
        }
        return type_map.get(ext, DocumentType.UNKNOWN)
    
    def parse(self, filepath: Path) -> Tuple[str, int, Dict[str, Any]]:
        """
        Parse a document and extract text.
        
        Args:
            filepath: Path to the document file
            
        Returns:
            Tuple of (extracted_text, page_count, metadata)
        """
        doc_type = self.get_document_type(filepath)
        
        if doc_type == DocumentType.PDF:
            return self._parse_pdf(filepath)
        elif doc_type == DocumentType.DOCX:
            return self._parse_docx(filepath)
        elif doc_type == DocumentType.TXT:
            return self._parse_txt(filepath)
        else:
            raise ValueError(f"Unsupported document type: {filepath.suffix}")
    
    def _parse_pdf(self, filepath: Path) -> Tuple[str, int, Dict[str, Any]]:
        """
        Parse PDF using multiple methods for best extraction.
        Uses pdfplumber for tables and layout-aware extraction,
        falls back to pypdf for simpler documents.
        """
        text_parts = []
        page_count = 0
        metadata = {}
        
        try:
            # Try pdfplumber first for better table extraction
            with pdfplumber.open(filepath) as pdf:
                page_count = len(pdf.pages)
                metadata["pdf_metadata"] = pdf.metadata or {}
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = ""
                    
                    # Extract tables first
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            table_text = self._format_table(table)
                            if table_text:
                                page_text += table_text + "\n\n"
                    
                    # Extract remaining text
                    main_text = page.extract_text() or ""
                    if main_text:
                        page_text += main_text
                    
                    if page_text.strip():
                        text_parts.append(f"[Page {page_num}]\n{page_text.strip()}")
                        
        except Exception as e:
            logger.warning(f"pdfplumber failed, falling back to pypdf: {e}")
            # Fallback to pypdf
            with open(filepath, "rb") as f:
                reader = pypdf.PdfReader(f)
                page_count = len(reader.pages)
                
                for page_num, page in enumerate(reader.pages, 1):
                    text = page.extract_text() or ""
                    if text.strip():
                        text_parts.append(f"[Page {page_num}]\n{text.strip()}")
        
        return "\n\n".join(text_parts), page_count, metadata
    
    def _format_table(self, table: List[List[str]]) -> str:
        """Format a table as readable text."""
        if not table:
            return ""
        
        formatted_rows = []
        for row in table:
            if row:
                # Clean and join cells
                cleaned = [str(cell).strip() if cell else "" for cell in row]
                if any(cleaned):  # Only add non-empty rows
                    formatted_rows.append(" | ".join(cleaned))
        
        return "\n".join(formatted_rows)
    
    def _parse_docx(self, filepath: Path) -> Tuple[str, int, Dict[str, Any]]:
        """Parse DOCX file extracting paragraphs and tables."""
        doc = DocxDocument(filepath)
        text_parts = []
        metadata = {
            "core_properties": {}
        }
        
        # Extract core properties if available
        try:
            props = doc.core_properties
            metadata["core_properties"] = {
                "author": props.author,
                "title": props.title,
                "subject": props.subject,
                "created": str(props.created) if props.created else None
            }
        except Exception:
            pass
        
        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)
        
        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(" | ".join(row_text))
            if table_text:
                text_parts.append("\n".join(table_text))
        
        # DOCX doesn't have traditional pages, estimate based on content
        full_text = "\n\n".join(text_parts)
        estimated_pages = max(1, len(full_text) // 3000)  # Rough estimate
        
        return full_text, estimated_pages, metadata
    
    def _parse_txt(self, filepath: Path) -> Tuple[str, int, Dict[str, Any]]:
        """Parse plain text file."""
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        
        # Estimate pages
        estimated_pages = max(1, len(text) // 3000)
        
        return text, estimated_pages, {}


class TextChunker:
    """
    Intelligent text chunking for logistics documents.
    Uses semantic boundaries (sections, paragraphs) when possible,
    with overlap to maintain context.
    """
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Patterns for logistics document sections
        self.section_patterns = [
            r'\n(?=SHIPPER|CONSIGNEE|CARRIER|PICKUP|DELIVERY|RATE|TOTAL|DESCRIPTION)',
            r'\n(?=Ship From|Ship To|Origin|Destination|Equipment)',
            r'\n(?=\d+\.\s+[A-Z])',  # Numbered sections
            r'\n(?=\[Page \d+\])',   # Page markers
            r'\n{2,}',               # Double newlines (paragraphs)
        ]
    
    def chunk(self, text: str, document_id: str) -> List[TextChunk]:
        """
        Split text into semantically meaningful chunks.
        
        Strategy:
        1. First try to split on document sections
        2. If sections are too large, split on paragraphs
        3. If still too large, split on sentences
        4. Apply overlap between chunks
        """
        if not text.strip():
            return []
        
        # Clean the text
        text = self._clean_text(text)
        
        # Get initial segments based on semantic boundaries
        segments = self._split_on_boundaries(text)
        
        # Process segments into appropriately sized chunks
        chunks = []
        current_chunk = ""
        current_start = 0
        char_position = 0
        
        for segment in segments:
            segment = segment.strip()
            if not segment:
                continue
            
            # If adding this segment exceeds chunk size
            if current_chunk and len(current_chunk) + len(segment) > self.chunk_size:
                # Save current chunk
                if current_chunk.strip():
                    chunks.append((current_chunk.strip(), current_start, char_position))
                
                # Start new chunk with overlap from previous
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + " " + segment if overlap_text else segment
                current_start = char_position - len(overlap_text) if overlap_text else char_position
            else:
                if current_chunk:
                    current_chunk += "\n" + segment
                else:
                    current_chunk = segment
                    current_start = char_position
            
            char_position += len(segment) + 1
        
        # Don't forget the last chunk
        if current_chunk.strip():
            chunks.append((current_chunk.strip(), current_start, char_position))
        
        # Handle any remaining oversized chunks
        final_chunks = []
        for chunk_text, start, end in chunks:
            if len(chunk_text) > self.chunk_size * 1.5:
                # Split oversized chunk on sentences
                sub_chunks = self._split_oversized_chunk(chunk_text, start)
                final_chunks.extend(sub_chunks)
            else:
                final_chunks.append((chunk_text, start, end))
        
        # Create TextChunk objects
        result = []
        for idx, (chunk_text, start, end) in enumerate(final_chunks):
            chunk_id = self._generate_chunk_id(document_id, idx, chunk_text)
            
            # Try to determine page number from page markers
            page_num = self._extract_page_number(chunk_text)
            
            result.append(TextChunk(
                content=chunk_text,
                chunk_id=chunk_id,
                document_id=document_id,
                chunk_index=idx,
                start_char=start,
                end_char=end,
                page_number=page_num,
                metadata={"char_count": len(chunk_text)}
            ))
        
        return result
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Normalize whitespace
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\t', ' ', text)
        text = re.sub(r' +', ' ', text)
        # Remove excessive newlines while preserving paragraph structure
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
    
    def _split_on_boundaries(self, text: str) -> List[str]:
        """Split text on semantic boundaries."""
        segments = [text]
        
        for pattern in self.section_patterns:
            new_segments = []
            for segment in segments:
                parts = re.split(pattern, segment)
                new_segments.extend(parts)
            segments = new_segments
        
        return [s.strip() for s in segments if s.strip()]
    
    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text from the end of a chunk."""
        if len(text) <= self.chunk_overlap:
            return text
        
        # Try to break at sentence boundary within overlap region
        overlap_region = text[-self.chunk_overlap * 2:]
        sentences = re.split(r'(?<=[.!?])\s+', overlap_region)
        
        if len(sentences) > 1:
            # Return last complete sentence(s) that fit in overlap
            result = ""
            for sent in reversed(sentences):
                if len(result) + len(sent) <= self.chunk_overlap:
                    result = sent + " " + result
                else:
                    break
            return result.strip()
        
        return text[-self.chunk_overlap:]
    
    def _split_oversized_chunk(self, text: str, start_pos: int) -> List[Tuple[str, int, int]]:
        """Split an oversized chunk into smaller pieces."""
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        chunks = []
        current = ""
        current_start = start_pos
        
        for sentence in sentences:
            if len(current) + len(sentence) > self.chunk_size:
                if current:
                    chunks.append((current.strip(), current_start, current_start + len(current)))
                current = sentence
                current_start = start_pos + len(text) - len(sentence)
            else:
                current += " " + sentence if current else sentence
        
        if current:
            chunks.append((current.strip(), current_start, current_start + len(current)))
        
        return chunks
    
    def _generate_chunk_id(self, document_id: str, index: int, content: str) -> str:
        """Generate a unique chunk ID."""
        hash_input = f"{document_id}_{index}_{content[:50]}"
        return hashlib.md5(hash_input.encode()).hexdigest()[:12]
    
    def _extract_page_number(self, text: str) -> Optional[int]:
        """Extract page number from page markers in text."""
        match = re.search(r'\[Page (\d+)\]', text)
        if match:
            return int(match.group(1))
        return None


class DocumentProcessor:
    """
    Main document processing pipeline.
    Combines parsing and chunking into a unified interface.
    """
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 100):
        self.parser = DocumentParser()
        self.chunker = TextChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    
    def process(self, filepath: Path) -> ParsedDocument:
        """
        Process a document file into a ParsedDocument.
        
        Args:
            filepath: Path to the document file
            
        Returns:
            ParsedDocument with extracted text and chunks
        """
        filepath = Path(filepath)
        
        if not filepath.exists():
            raise FileNotFoundError(f"Document not found: {filepath}")
        
        # Generate document ID from filename and content hash
        document_id = self._generate_document_id(filepath)
        
        # Parse the document
        doc_type = self.parser.get_document_type(filepath)
        raw_text, page_count, metadata = self.parser.parse(filepath)
        
        # Chunk the text
        chunks = self.chunker.chunk(raw_text, document_id)
        
        logger.info(f"Processed document: {filepath.name} - {len(chunks)} chunks, {page_count} pages")
        
        return ParsedDocument(
            document_id=document_id,
            filename=filepath.name,
            document_type=doc_type,
            raw_text=raw_text,
            chunks=chunks,
            page_count=page_count,
            char_count=len(raw_text),
            metadata=metadata
        )
    
    def _generate_document_id(self, filepath: Path) -> str:
        """Generate a unique document ID."""
        with open(filepath, "rb") as f:
            content_hash = hashlib.md5(f.read()).hexdigest()[:8]
        return f"{filepath.stem}_{content_hash}"
